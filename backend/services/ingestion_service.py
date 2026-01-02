"""
Data ingestion service for Git repositories, Confluence, and documents.
"""
from typing import List, Dict, Any, Optional, Tuple
import os
import tempfile
import shutil
from pathlib import Path
import asyncio
import aiofiles
import structlog
from git import Repo, GitCommandError
import requests
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document as DocxDocument

# Advanced PDF processing with Unstructured.io
try:
    from unstructured.partition.pdf import partition_pdf
    from unstructured.documents.elements import (
        Title, NarrativeText, ListItem, Table, Image, 
        Header, Footer, PageBreak, FigureCaption
    )
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    UNSTRUCTURED_AVAILABLE = False

from sqlalchemy import func

from core.config import get_settings
from core.database import get_db, DataSource, Document, DocumentChunk
from services.vector_service import VectorService

logger = structlog.get_logger()
settings = get_settings()

class TextChunker:
    """Service for chunking text into smaller pieces with line tracking."""
    
    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
    
    def _build_line_index(self, text: str) -> List[int]:
        """Build an index mapping character positions to line numbers."""
        line_starts = [0]
        for i, char in enumerate(text):
            if char == '\n':
                line_starts.append(i + 1)
        return line_starts
    
    def _char_to_line(self, char_pos: int, line_starts: List[int]) -> int:
        """Convert character position to line number (1-indexed)."""
        for i, start in enumerate(line_starts):
            if char_pos < start:
                return i  # 1-indexed
        return len(line_starts)
    
    def chunk_text(self, text: str, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks with line number tracking."""
        if not text.strip():
            return []
        
        # Build line index for accurate line number tracking
        line_starts = self._build_line_index(text)
        
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk_text = text[start:end]
            
            # Try to break at sentence boundaries
            if end < len(text):
                last_period = chunk_text.rfind('.')
                last_newline = chunk_text.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > self.chunk_size // 2:
                    chunk_text = text[start:start + break_point + 1]
                    end = start + break_point + 1
            
            # Calculate line numbers for this chunk
            start_line = self._char_to_line(start, line_starts)
            end_line = self._char_to_line(min(end, len(text) - 1), line_starts)
            
            chunk_metadata = metadata.copy() if metadata else {}
            chunk_metadata.update({
                'chunk_index': chunk_index,
                'start_char': start,
                'end_char': end,
                'start_line': start_line,
                'end_line': end_line
            })
            
            chunks.append({
                'content': chunk_text.strip(),
                'metadata': chunk_metadata
            })
            
            start = end - self.chunk_overlap
            chunk_index += 1
        
        return chunks

class GitIngestionService:
    """Service for ingesting Git repositories."""
    
    def __init__(self):
        self.clone_dir = Path(settings.GIT_CLONE_DIR)
        self.clone_dir.mkdir(parents=True, exist_ok=True)
        self.chunker = TextChunker()
    
    # Language filter presets
    LANGUAGE_FILTERS = {
        'all': {'.py', '.js', '.ts', '.tsx', '.jsx', '.java', '.cpp', '.c', '.h', '.hpp', '.cc', '.cxx',
                '.md', '.txt', '.rst', '.json', '.yaml', '.yml', '.go', '.rs', '.rb', '.php', '.cs', '.swift'},
        'c_cpp': {'.c', '.h', '.cpp', '.cc', '.cxx', '.hpp', '.hxx', '.hh'},
        'python': {'.py', '.pyi', '.pyx'},
        'javascript': {'.js', '.jsx', '.ts', '.tsx', '.mjs', '.cjs'},
        'java': {'.java'},
        'docs': {'.md', '.txt', '.rst', '.adoc', '.org'},
        'go': {'.go'},
        'rust': {'.rs'},
    }
    
    async def ingest_repository(
        self, 
        repo_url: str, 
        workspace_id: int, 
        branch: str = "main",
        language_filter: Optional[str] = None,
        max_depth: Optional[int] = None,
        username: Optional[str] = None,
        token: Optional[str] = None
    ) -> Dict[str, Any]:
        """Clone and ingest a Git repository.
        
        Supports multiple authentication methods:
        - No auth: Public repositories
        - Bearer token: For Bitbucket Server with Basic Auth disabled
        - Basic auth: username:token in URL (for GitHub, GitLab, etc.)
        """
        try:
            logger.info(f"Starting Git ingestion for {repo_url} (language={language_filter}, max_depth={max_depth})")
            
            # Determine which extensions to include
            if language_filter and language_filter in self.LANGUAGE_FILTERS:
                supported_extensions = self.LANGUAGE_FILTERS[language_filter]
            else:
                # Default: all supported extensions
                supported_extensions = self.LANGUAGE_FILTERS['all']
            
            # Create temporary directory for cloning
            with tempfile.TemporaryDirectory() as temp_dir:
                repo_path = Path(temp_dir) / "repo"
                base_depth = len(repo_path.parts)
                
                # Clone repository with authentication
                try:
                    import subprocess
                    
                    if token:
                        # Use Bearer token authentication (for Bitbucket Server with Basic Auth disabled)
                        # Use subprocess directly since GitPython blocks -c option
                        logger.info(f"Using Bearer token authentication for {repo_url}")
                        cmd = [
                            'git', '-c', f'http.extraHeader=Authorization: Bearer {token}',
                            'clone', '-v', f'--branch={branch}', '--depth=1',
                            '--', repo_url, str(repo_path)
                        ]
                        result = subprocess.run(cmd, capture_output=True, text=True)
                        if result.returncode != 0:
                            raise GitCommandError(cmd, result.returncode, result.stderr)
                        repo = Repo(repo_path)
                    elif username and token:
                        # Fallback: embed credentials in URL (for GitHub, GitLab, etc.)
                        # URL-encode the token in case it contains special characters
                        from urllib.parse import quote
                        encoded_token = quote(token, safe='')
                        if '://' in repo_url:
                            protocol, rest = repo_url.split('://', 1)
                            repo_url = f"{protocol}://{username}:{encoded_token}@{rest}"
                            logger.info(f"Using Basic auth (URL-embedded) for repository")
                        repo = Repo.clone_from(repo_url, repo_path, branch=branch, depth=1)
                    else:
                        # No authentication - public repository
                        repo = Repo.clone_from(repo_url, repo_path, branch=branch, depth=1)
                    
                    logger.info(f"Cloned repository to {repo_path}")
                except GitCommandError as e:
                    logger.error(f"Failed to clone repository: {e}")
                    return {"success": False, "error": str(e)}
                
                # Process files
                documents = []
                
                for file_path in repo_path.rglob('*'):
                    if not file_path.is_file():
                        continue
                    
                    # Check extension
                    if file_path.suffix.lower() not in supported_extensions:
                        continue
                    
                    # Check depth
                    if max_depth is not None:
                        file_depth = len(file_path.parts) - base_depth
                        if file_depth > max_depth:
                            continue
                    
                    try:
                        content = await self._read_file(file_path)
                        if content:
                            relative_path = file_path.relative_to(repo_path)
                            documents.append({
                                'title': str(relative_path),
                                'content': content,
                                'file_path': str(relative_path),
                                'file_type': file_path.suffix.lower(),
                                'source_type': 'git',
                                'workspace_id': workspace_id,
                                'metadata': {
                                    'repo_url': repo_url,
                                    'branch': branch,
                                    'file_size': file_path.stat().st_size,
                                    'language_filter': language_filter
                                }
                            })
                    except Exception as e:
                        logger.warning(f"Failed to read file {file_path}: {e}")
                
                logger.info(f"Processed {len(documents)} files from repository")
                return {"success": True, "documents": documents}
                
        except Exception as e:
            logger.error(f"Error ingesting Git repository: {e}")
            return {"success": False, "error": str(e)}
    
    async def _read_file(self, file_path: Path) -> Optional[str]:
        """Read file content safely."""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = await f.read()
                return content
        except Exception as e:
            logger.warning(f"Failed to read {file_path}: {e}")
            return None

class ConfluenceIngestionService:
    """Service for ingesting Confluence pages."""
    
    def __init__(self):
        self.chunker = TextChunker()
    
    async def ingest_space(
        self, 
        space_key: str, 
        workspace_id: int, 
        base_url: str = None, 
        username: str = None, 
        api_token: str = None,
        page_ids: List[str] = None,
        max_depth: int = None,
        include_children: bool = True
    ) -> Dict[str, Any]:
        """
        Ingest pages from a Confluence space.
        
        Args:
            space_key: Confluence space key
            workspace_id: Target workspace ID
            base_url: Confluence base URL
            username: Confluence username/email
            api_token: Confluence API token
            page_ids: Optional list of specific page IDs to ingest
            max_depth: Max depth for child pages (None = all)
            include_children: Whether to include child pages
        """
        try:
            base_url = base_url or settings.CONFLUENCE_BASE_URL if hasattr(settings, 'CONFLUENCE_BASE_URL') else None
            username = username or settings.CONFLUENCE_USERNAME if hasattr(settings, 'CONFLUENCE_USERNAME') else None
            api_token = api_token or settings.CONFLUENCE_API_TOKEN if hasattr(settings, 'CONFLUENCE_API_TOKEN') else None
            
            if not all([base_url, username, api_token]):
                return {"success": False, "error": "Missing Confluence credentials"}
            
            # Ensure base_url doesn't have trailing slash
            base_url = base_url.rstrip('/')
            
            logger.info(f"Starting Confluence ingestion for space {space_key} (page_ids={page_ids}, max_depth={max_depth})")
            
            # Get pages from space
            if page_ids:
                # Fetch specific pages
                pages = await self._get_specific_pages(base_url, page_ids, username, api_token)
            else:
                # Fetch all pages from space
                pages = await self._get_space_pages(base_url, space_key, username, api_token, max_depth, include_children)
            
            documents = []
            for page in pages:
                try:
                    content = await self._get_page_content(base_url, page['id'], username, api_token)
                    if content:
                        # Extract additional metadata
                        page_version = page.get('version', {}).get('number', 1)
                        page_created = page.get('history', {}).get('createdDate', '')
                        page_updated = page.get('version', {}).get('when', '')
                        page_creator = page.get('history', {}).get('createdBy', {}).get('displayName', 'Unknown')
                        page_modifier = page.get('version', {}).get('by', {}).get('displayName', 'Unknown')
                        ancestors = [a.get('title', '') for a in page.get('ancestors', [])]
                        
                        documents.append({
                            'title': page['title'],
                            'content': content,
                            'file_path': f"confluence/{space_key}/{page['id']}",
                            'file_type': 'confluence',
                            'source_type': 'confluence',
                            'workspace_id': workspace_id,
                            'metadata': {
                                'space_key': space_key,
                                'page_id': page['id'],
                                'page_title': page['title'],
                                'page_version': page_version,
                                'created': page_created,
                                'updated': page_updated,
                                'creator': page_creator,
                                'last_modifier': page_modifier,
                                'ancestors': ancestors,
                                'page_url': f"{base_url}/pages/viewpage.action?pageId={page['id']}"
                            }
                        })
                except Exception as e:
                    logger.warning(f"Failed to process page {page['id']}: {e}")
            
            logger.info(f"Processed {len(documents)} pages from Confluence space")
            return {"success": True, "documents": documents}
            
        except Exception as e:
            logger.error(f"Error ingesting Confluence space: {e}")
            return {"success": False, "error": str(e)}
    
    async def _get_space_pages(
        self, 
        base_url: str, 
        space_key: str, 
        username: str, 
        api_token: str,
        max_depth: int = None,
        include_children: bool = True
    ) -> List[Dict[str, Any]]:
        """Get pages from a Confluence space with optional depth limiting."""
        pages = []
        start = 0
        limit = 50
        
        while True:
            url = f"{base_url}/rest/api/content"
            params = {
                'spaceKey': space_key,
                'type': 'page',
                'status': 'current',
                'start': start,
                'limit': limit,
                'expand': 'version,history,ancestors'
            }
            
            response = requests.get(url, params=params, auth=(username, api_token))
            response.raise_for_status()
            
            data = response.json()
            
            for page in data['results']:
                # Filter by depth if max_depth is specified
                if max_depth is not None:
                    ancestors = page.get('ancestors', [])
                    if len(ancestors) > max_depth:
                        continue
                
                pages.append(page)
            
            if len(data['results']) < limit:
                break
            
            start += limit
        
        return pages
    
    async def _get_specific_pages(
        self, 
        base_url: str, 
        page_ids: List[str], 
        username: str, 
        api_token: str
    ) -> List[Dict[str, Any]]:
        """Get specific pages by their IDs."""
        pages = []
        
        for page_id in page_ids:
            try:
                url = f"{base_url}/rest/api/content/{page_id}"
                params = {'expand': 'version,history,ancestors'}
                
                response = requests.get(url, params=params, auth=(username, api_token))
                response.raise_for_status()
                
                pages.append(response.json())
            except Exception as e:
                logger.warning(f"Failed to fetch page {page_id}: {e}")
        
        return pages
    
    async def _get_page_content(self, base_url: str, page_id: str, 
                               username: str, api_token: str) -> Optional[str]:
        """Get content of a specific Confluence page."""
        try:
            url = f"{base_url}/rest/api/content/{page_id}"
            params = {'expand': 'body.storage'}
            
            response = requests.get(url, params=params, auth=(username, api_token))
            response.raise_for_status()
            
            data = response.json()
            html_content = data['body']['storage']['value']
            
            # Convert HTML to plain text
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text(separator='\n', strip=True)
            
            return text_content
            
        except Exception as e:
            logger.warning(f"Failed to get content for page {page_id}: {e}")
            return None

class JiraIngestionService:
    """Service for ingesting JIRA issues."""
    
    def __init__(self):
        self.chunker = TextChunker()
    
    async def ingest_project(
        self, 
        project_key: str, 
        workspace_id: int,
        base_url: str = None,
        username: str = None,
        api_token: str = None,
        issue_types: List[str] = None,
        specific_tickets: List[str] = None,
        max_results: int = None
    ) -> Dict[str, Any]:
        """
        Ingest issues from a JIRA project.
        
        Args:
            project_key: JIRA project key (e.g., 'PROJ')
            workspace_id: Target workspace ID
            base_url: JIRA base URL (e.g., 'https://company.atlassian.net')
            username: JIRA username/email
            api_token: JIRA API token
            issue_types: Filter by issue types ['Story', 'Epic', 'Bug', 'Task'] or None for all
            specific_tickets: List of specific ticket keys to ingest (e.g., ['PROJ-123', 'PROJ-456'])
            max_results: Maximum number of issues to fetch (None for all)
        """
        try:
            base_url = base_url or settings.JIRA_BASE_URL if hasattr(settings, 'JIRA_BASE_URL') else None
            username = username or settings.JIRA_USERNAME if hasattr(settings, 'JIRA_USERNAME') else None
            api_token = api_token or settings.JIRA_API_TOKEN if hasattr(settings, 'JIRA_API_TOKEN') else None
            
            if not all([base_url, username, api_token]):
                return {"success": False, "error": "Missing JIRA credentials"}
            
            # Ensure base_url doesn't have trailing slash
            base_url = base_url.rstrip('/')
            
            logger.info(f"Starting JIRA ingestion for project {project_key} (types={issue_types}, specific={specific_tickets})")
            
            # Get issues from project
            issues = await self._get_project_issues(
                base_url, project_key, username, api_token,
                issue_types=issue_types,
                specific_tickets=specific_tickets,
                max_results=max_results
            )
            
            documents = []
            for issue in issues:
                try:
                    content = self._format_issue_content(issue)
                    if content:
                        issue_key = issue['key']
                        issue_type = issue['fields'].get('issuetype', {}).get('name', 'Unknown')
                        summary = issue['fields'].get('summary', 'No summary')
                        status = issue['fields'].get('status', {}).get('name', 'Unknown')
                        priority = issue['fields'].get('priority', {}).get('name', 'None')
                        assignee = issue['fields'].get('assignee', {})
                        assignee_name = assignee.get('displayName', 'Unassigned') if assignee else 'Unassigned'
                        reporter = issue['fields'].get('reporter', {})
                        reporter_name = reporter.get('displayName', 'Unknown') if reporter else 'Unknown'
                        created = issue['fields'].get('created', '')
                        updated = issue['fields'].get('updated', '')
                        labels = issue['fields'].get('labels', [])
                        components = [c.get('name', '') for c in issue['fields'].get('components', [])]
                        
                        documents.append({
                            'title': f"[{issue_key}] {summary}",
                            'content': content,
                            'file_path': f"jira/{project_key}/{issue_key}",
                            'file_type': 'jira',
                            'source_type': 'jira',
                            'workspace_id': workspace_id,
                            'metadata': {
                                'project_key': project_key,
                                'issue_key': issue_key,
                                'issue_type': issue_type,
                                'status': status,
                                'priority': priority,
                                'assignee': assignee_name,
                                'reporter': reporter_name,
                                'created': created,
                                'updated': updated,
                                'labels': labels,
                                'components': components,
                                'issue_url': f"{base_url}/browse/{issue_key}"
                            }
                        })
                except Exception as e:
                    logger.warning(f"Failed to process issue {issue.get('key', 'unknown')}: {e}")
            
            logger.info(f"Processed {len(documents)} issues from JIRA project {project_key}")
            return {"success": True, "documents": documents}
            
        except Exception as e:
            logger.error(f"Error ingesting JIRA project: {e}")
            return {"success": False, "error": str(e)}
    
    async def _get_project_issues(
        self, 
        base_url: str, 
        project_key: str,
        username: str, 
        api_token: str,
        issue_types: List[str] = None,
        specific_tickets: List[str] = None,
        max_results: int = None
    ) -> List[Dict[str, Any]]:
        """Get issues from a JIRA project with optional filtering."""
        issues = []
        start_at = 0
        page_size = 50
        
        # Build JQL query
        if specific_tickets:
            # Fetch specific tickets
            jql = f"key in ({','.join(specific_tickets)})"
        else:
            jql_parts = [f"project = {project_key}"]
            if issue_types:
                type_filter = ','.join([f'"{t}"' for t in issue_types])
                jql_parts.append(f"issuetype in ({type_filter})")
            jql = ' AND '.join(jql_parts)
            jql += " ORDER BY updated DESC"
        
        logger.info(f"JIRA JQL: {jql}")
        
        while True:
            url = f"{base_url}/rest/api/3/search"
            params = {
                'jql': jql,
                'startAt': start_at,
                'maxResults': page_size,
                'fields': 'summary,description,issuetype,status,priority,assignee,reporter,created,updated,labels,components,comment'
            }
            
            response = requests.get(url, params=params, auth=(username, api_token))
            response.raise_for_status()
            
            data = response.json()
            issues.extend(data['issues'])
            
            # Check if we've reached the limit or end
            if max_results and len(issues) >= max_results:
                issues = issues[:max_results]
                break
            
            if start_at + page_size >= data['total']:
                break
            
            start_at += page_size
        
        return issues
    
    def _format_issue_content(self, issue: Dict[str, Any]) -> str:
        """Format JIRA issue into readable text content."""
        fields = issue['fields']
        
        parts = []
        
        # Header
        parts.append(f"# {issue['key']}: {fields.get('summary', 'No summary')}")
        parts.append("")
        
        # Metadata
        parts.append("## Issue Details")
        parts.append(f"- **Type:** {fields.get('issuetype', {}).get('name', 'Unknown')}")
        parts.append(f"- **Status:** {fields.get('status', {}).get('name', 'Unknown')}")
        parts.append(f"- **Priority:** {fields.get('priority', {}).get('name', 'None')}")
        
        assignee = fields.get('assignee')
        parts.append(f"- **Assignee:** {assignee.get('displayName', 'Unassigned') if assignee else 'Unassigned'}")
        
        reporter = fields.get('reporter')
        parts.append(f"- **Reporter:** {reporter.get('displayName', 'Unknown') if reporter else 'Unknown'}")
        
        if fields.get('labels'):
            parts.append(f"- **Labels:** {', '.join(fields['labels'])}")
        
        components = fields.get('components', [])
        if components:
            parts.append(f"- **Components:** {', '.join([c.get('name', '') for c in components])}")
        
        parts.append("")
        
        # Description
        description = fields.get('description')
        if description:
            parts.append("## Description")
            # Handle Atlassian Document Format (ADF)
            if isinstance(description, dict):
                parts.append(self._parse_adf_content(description))
            else:
                parts.append(str(description))
            parts.append("")
        
        # Comments
        comments = fields.get('comment', {}).get('comments', [])
        if comments:
            parts.append("## Comments")
            for comment in comments:
                author = comment.get('author', {}).get('displayName', 'Unknown')
                created = comment.get('created', '')[:10]  # Just date
                body = comment.get('body', '')
                
                # Handle ADF format for comments
                if isinstance(body, dict):
                    body = self._parse_adf_content(body)
                
                parts.append(f"### {author} ({created})")
                parts.append(body)
                parts.append("")
        
        return '\n'.join(parts)
    
    def _parse_adf_content(self, adf: Dict[str, Any]) -> str:
        """Parse Atlassian Document Format to plain text."""
        if not adf or not isinstance(adf, dict):
            return ""
        
        content_parts = []
        
        def extract_text(node):
            if isinstance(node, str):
                return node
            if isinstance(node, dict):
                if node.get('type') == 'text':
                    return node.get('text', '')
                if 'content' in node:
                    return ''.join(extract_text(child) for child in node['content'])
            if isinstance(node, list):
                return ''.join(extract_text(item) for item in node)
            return ''
        
        for block in adf.get('content', []):
            block_type = block.get('type', '')
            text = extract_text(block)
            
            if block_type == 'heading':
                level = block.get('attrs', {}).get('level', 1)
                content_parts.append(f"{'#' * level} {text}")
            elif block_type == 'bulletList':
                for item in block.get('content', []):
                    item_text = extract_text(item)
                    content_parts.append(f"- {item_text}")
            elif block_type == 'orderedList':
                for i, item in enumerate(block.get('content', []), 1):
                    item_text = extract_text(item)
                    content_parts.append(f"{i}. {item_text}")
            elif block_type == 'codeBlock':
                lang = block.get('attrs', {}).get('language', '')
                content_parts.append(f"```{lang}")
                content_parts.append(text)
                content_parts.append("```")
            else:
                if text.strip():
                    content_parts.append(text)
            
            content_parts.append("")
        
        return '\n'.join(content_parts)


class PDFExtractionStrategy:
    """Enum-like class for PDF extraction strategies."""
    UNSTRUCTURED_HI_RES = "unstructured_hi_res"  # Best quality, slower
    UNSTRUCTURED_FAST = "unstructured_fast"      # Faster, good quality
    PYPDF2_FALLBACK = "pypdf2_fallback"          # Basic fallback


class DocumentIngestionService:
    """Service for ingesting PDF and Word documents with advanced parsing."""
    
    def __init__(self, pdf_strategy: str = None):
        self.chunker = TextChunker()
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Set PDF extraction strategy
        if pdf_strategy:
            self.pdf_strategy = pdf_strategy
        elif UNSTRUCTURED_AVAILABLE:
            # Use FAST strategy by default for better performance
            self.pdf_strategy = PDFExtractionStrategy.UNSTRUCTURED_FAST
        else:
            self.pdf_strategy = PDFExtractionStrategy.PYPDF2_FALLBACK
            logger.warning("Unstructured.io not available, falling back to PyPDF2")
    
    async def ingest_document(self, file_path: str, workspace_id: int, 
                             original_filename: str = None) -> Dict[str, Any]:
        """Ingest a document file."""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {"success": False, "error": "File not found"}
            
            logger.info(f"Starting document ingestion for {file_path}")
            
            # Extract text based on file type
            extraction_metadata = {}
            if file_path.suffix.lower() == '.pdf':
                content, extraction_metadata = await self._extract_pdf_text(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                content = await self._extract_docx_text(file_path)
            else:
                # Try to read as text file
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = await f.read()
            
            if not content:
                return {"success": False, "error": "Could not extract text from document"}
            
            # Build metadata with extraction info
            doc_metadata = {
                'original_filename': original_filename,
                'file_size': file_path.stat().st_size
            }
            if extraction_metadata:
                doc_metadata['extraction'] = extraction_metadata
            
            documents = [{
                'title': original_filename or file_path.name,
                'content': content,
                'file_path': str(file_path),
                'file_type': file_path.suffix.lower(),
                'source_type': 'document',
                'workspace_id': workspace_id,
                'metadata': doc_metadata
            }]
            
            logger.info(f"Extracted text from document: {len(content)} characters")
            return {"success": True, "documents": documents, "extraction_metadata": extraction_metadata}
            
        except Exception as e:
            logger.error(f"Error ingesting document: {e}")
            return {"success": False, "error": str(e)}
    
    async def _extract_pdf_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF file using the configured strategy.
        Returns tuple of (text, extraction_metadata).
        """
        extraction_metadata = {
            "strategy_used": self.pdf_strategy,
            "tables_found": 0,
            "images_found": 0,
            "pages_processed": 0,
            "extraction_quality": "unknown"
        }
        
        # Try Unstructured.io first if available and configured
        if UNSTRUCTURED_AVAILABLE and self.pdf_strategy in [
            PDFExtractionStrategy.UNSTRUCTURED_HI_RES,
            PDFExtractionStrategy.UNSTRUCTURED_FAST
        ]:
            text, metadata = await self._extract_with_unstructured(file_path)
            if text:
                extraction_metadata.update(metadata)
                return text, extraction_metadata
            else:
                logger.warning(f"Unstructured extraction failed, falling back to PyPDF2")
                extraction_metadata["strategy_used"] = PDFExtractionStrategy.PYPDF2_FALLBACK
        
        # Fallback to PyPDF2
        text, metadata = await self._extract_with_pypdf2(file_path)
        extraction_metadata.update(metadata)
        return text, extraction_metadata
    
    async def _extract_with_unstructured(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text using Unstructured.io with advanced parsing.
        Handles tables, images, and complex layouts.
        """
        metadata = {
            "tables_found": 0,
            "images_found": 0,
            "pages_processed": 0,
            "extraction_quality": "high"
        }
        
        try:
            # Configure extraction strategy
            strategy = "hi_res" if self.pdf_strategy == PDFExtractionStrategy.UNSTRUCTURED_HI_RES else "fast"
            
            logger.info(f"Extracting PDF with Unstructured.io ({strategy} strategy): {file_path}")
            
            # Partition the PDF
            elements = partition_pdf(
                filename=str(file_path),
                strategy=strategy,
                infer_table_structure=True,  # Extract tables as structured data
                include_page_breaks=True,
                extract_images_in_pdf=False,  # Set to True if you want image extraction
            )
            
            # Process elements and build structured text
            text_parts = []
            current_page = 1
            
            for element in elements:
                # Track page breaks
                if isinstance(element, PageBreak):
                    current_page += 1
                    text_parts.append(f"\n--- Page {current_page} ---\n")
                    continue
                
                # Handle different element types
                if isinstance(element, Title):
                    text_parts.append(f"\n## {element.text}\n")
                elif isinstance(element, Header):
                    text_parts.append(f"\n### {element.text}\n")
                elif isinstance(element, Table):
                    metadata["tables_found"] += 1
                    # Convert table to markdown format
                    table_text = self._table_to_markdown(element)
                    text_parts.append(f"\n{table_text}\n")
                elif isinstance(element, Image):
                    metadata["images_found"] += 1
                    text_parts.append(f"\n[Image: {element.text if element.text else 'Figure'}]\n")
                elif isinstance(element, FigureCaption):
                    text_parts.append(f"\n*Figure: {element.text}*\n")
                elif isinstance(element, ListItem):
                    text_parts.append(f"• {element.text}\n")
                elif isinstance(element, (NarrativeText, Footer)):
                    text_parts.append(f"{element.text}\n")
                else:
                    # Default handling for other element types
                    if hasattr(element, 'text') and element.text:
                        text_parts.append(f"{element.text}\n")
            
            metadata["pages_processed"] = current_page
            
            full_text = "\n".join(text_parts)
            logger.info(f"Unstructured extraction complete: {len(full_text)} chars, "
                       f"{metadata['tables_found']} tables, {metadata['images_found']} images")
            
            return full_text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error in Unstructured PDF extraction: {e}")
            metadata["extraction_quality"] = "failed"
            return "", metadata
    
    def _table_to_markdown(self, table_element) -> str:
        """
        Convert an Unstructured Table element to markdown format.
        """
        try:
            # Check if table has HTML representation
            if hasattr(table_element, 'metadata') and hasattr(table_element.metadata, 'text_as_html'):
                html = table_element.metadata.text_as_html
                if html:
                    return self._html_table_to_markdown(html)
            
            # Fallback to plain text representation
            return f"[Table]\n{table_element.text}\n[/Table]"
            
        except Exception as e:
            logger.warning(f"Error converting table to markdown: {e}")
            return f"[Table]\n{table_element.text if hasattr(table_element, 'text') else 'Unable to parse table'}\n[/Table]"
    
    def _html_table_to_markdown(self, html: str) -> str:
        """
        Convert HTML table to markdown format.
        """
        try:
            soup = BeautifulSoup(html, 'html.parser')
            table = soup.find('table')
            if not table:
                return html
            
            rows = table.find_all('tr')
            if not rows:
                return html
            
            markdown_rows = []
            
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                cell_texts = [cell.get_text(strip=True) for cell in cells]
                markdown_rows.append("| " + " | ".join(cell_texts) + " |")
                
                # Add header separator after first row
                if i == 0:
                    separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                    markdown_rows.append(separator)
            
            return "\n".join(markdown_rows)
            
        except Exception as e:
            logger.warning(f"Error converting HTML table to markdown: {e}")
            return html
    
    async def _extract_with_pypdf2(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF using PyPDF2 (basic fallback).
        """
        metadata = {
            "tables_found": 0,
            "images_found": 0,
            "pages_processed": 0,
            "extraction_quality": "basic"
        }
        
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages_processed"] = len(pdf_reader.pages)
                
                for i, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {i + 1} ---\n{page_text}")
            
            return "\n\n".join(text_parts).strip(), metadata
            
        except Exception as e:
            logger.error(f"Error extracting PDF text with PyPDF2: {e}")
            metadata["extraction_quality"] = "failed"
            return "", metadata
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from Word document."""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

class IngestionOrchestrator:
    """Main ingestion service that coordinates all ingestion types."""
    
    def __init__(self):
        self.git_service = GitIngestionService()
        self.confluence_service = ConfluenceIngestionService()
        self.jira_service = JiraIngestionService()
        self.document_service = DocumentIngestionService()
        self.vector_service = VectorService()
        self.chunker = TextChunker()
    
    async def ingest_data_source(self, data_source_id: int, progress_callback=None) -> Dict[str, Any]:
        """Ingest data from a data source."""
        try:
            # Get data source from database
            db = next(get_db())
            data_source = db.query(DataSource).filter(DataSource.id == data_source_id).first()
            
            if not data_source:
                return {"success": False, "error": "Data source not found"}
            
            # Update status to processing
            data_source.status = "processing"
            db.commit()
            
            # Ingest based on source type
            logger.info(f"[STAGE 1/4] Starting ingestion for data source {data_source_id} (type: {data_source.source_type})")
            if progress_callback:
                progress_callback(data_source_id, "Cloning Repository", 1, 4, 0, 1, "Fetching files from source...")
            
            if data_source.source_type == "git":
                config = data_source.config or {}
                logger.info(f"[GIT] Cloning repository: {data_source.source_url}")
                result = await self.git_service.ingest_repository(
                    data_source.source_url, 
                    data_source.workspace_id,
                    branch=config.get('branch', 'main'),
                    language_filter=config.get('language_filter'),
                    max_depth=config.get('max_depth'),
                    username=config.get('username'),
                    token=config.get('token')
                )
            elif data_source.source_type == "confluence":
                config = data_source.config or {}
                result = await self.confluence_service.ingest_space(
                    config.get('space_key'),
                    data_source.workspace_id,
                    config.get('base_url'),
                    config.get('username'),
                    config.get('api_token'),
                    page_ids=config.get('page_ids'),
                    max_depth=config.get('max_depth'),
                    include_children=config.get('include_children', True)
                )
            elif data_source.source_type == "jira":
                config = data_source.config or {}
                logger.info(f"[JIRA] Fetching issues from project: {config.get('project_key')}")
                result = await self.jira_service.ingest_project(
                    config.get('project_key'),
                    data_source.workspace_id,
                    config.get('base_url'),
                    config.get('username'),
                    config.get('api_token'),
                    issue_types=config.get('issue_types'),
                    specific_tickets=config.get('specific_tickets'),
                    max_results=config.get('max_results')
                )
            elif data_source.source_type == "document":
                result = await self.document_service.ingest_document(
                    data_source.source_url,  # File path for documents
                    data_source.workspace_id,
                    data_source.config.get('original_filename') if data_source.config else None
                )
            else:
                result = {"success": False, "error": f"Unsupported source type: {data_source.source_type}"}
            
            if result["success"]:
                # Process and store documents
                documents = result["documents"]
                logger.info(f"[GIT] Clone complete. Found {len(documents)} files to process.")
                if progress_callback:
                    progress_callback(data_source_id, "Processing Documents", 2, 4, 0, len(documents), f"Processing {len(documents)} files...")
                
                await self._process_documents(documents, data_source_id, db, progress_callback)
                
                # Update status to completed
                data_source.status = "completed"
                data_source.last_ingested = func.now()
                db.commit()
                
                if progress_callback:
                    progress_callback(data_source_id, "Complete", 4, 4, len(documents), len(documents), "Ingestion complete!")
                
                logger.info(f"Successfully ingested {len(documents)} documents from data source {data_source_id}")
                return {"success": True, "documents_count": len(documents)}
            else:
                # Update status to failed
                data_source.status = "failed"
                db.commit()
                return result
                
        except Exception as e:
            logger.error(f"Error in ingestion orchestrator: {e}")
            return {"success": False, "error": str(e)}
    
    async def _process_documents(self, documents: List[Dict[str, Any]], 
                               data_source_id: int, db,
                               progress_callback=None) -> None:
        """Process documents: chunk, embed, and store."""
        try:
            all_chunks = []
            total_docs = len(documents)
            
            logger.info(f"[STAGE 2/4] Starting document processing for {total_docs} documents")
            if progress_callback:
                progress_callback(data_source_id, "Processing Documents", 2, 4, 0, total_docs, "Chunking files...")
            
            for idx, doc_data in enumerate(documents):
                doc_title = doc_data['title'][:50] if doc_data['title'] else 'Unknown'
                logger.info(f"[DOC {idx+1}/{total_docs}] Processing: {doc_title}")
                
                # Update progress every 5 documents
                if progress_callback and (idx % 5 == 0 or idx == total_docs - 1):
                    progress_callback(data_source_id, "Processing Documents", 2, 4, idx + 1, total_docs, f"Processing: {doc_title}")
                
                # Create document record
                document = Document(
                    data_source_id=data_source_id,
                    title=doc_data['title'],
                    content=doc_data['content'],
                    file_path=doc_data.get('file_path'),
                    file_type=doc_data.get('file_type'),
                    doc_metadata=doc_data.get('metadata', {})
                )
                db.add(document)
                db.flush()  # Get the ID
                
                # Chunk the document
                logger.debug(f"[DOC {idx+1}/{total_docs}] Chunking document...")
                chunks = self.chunker.chunk_text(doc_data['content'], doc_data.get('metadata', {}))
                logger.info(f"[DOC {idx+1}/{total_docs}] Created {len(chunks)} chunks")
                
                for chunk_data in chunks:
                    chunk = DocumentChunk(
                        document_id=document.id,
                        chunk_index=chunk_data['metadata']['chunk_index'],
                        content=chunk_data['content'],
                        chunk_metadata=chunk_data['metadata']
                    )
                    db.add(chunk)
                    db.flush()
                    
                    # Prepare for vector storage
                    all_chunks.append({
                        'id': f"chunk_{chunk.id}",
                        'content': chunk_data['content'],
                        'title': doc_data['title'],
                        'source': doc_data.get('file_path', ''),
                        'workspace_id': doc_data['workspace_id'],
                        'document_id': document.id,
                        'chunk_id': chunk.id,
                        **chunk_data['metadata']
                    })
                
                # Report progress every 10 documents
                if (idx + 1) % 10 == 0:
                    logger.info(f"[PROGRESS] Processed {idx+1}/{total_docs} documents ({len(all_chunks)} chunks so far)")
            
            logger.info(f"[STAGE 3/4] Document processing complete. Total chunks: {len(all_chunks)}")
            if progress_callback:
                progress_callback(data_source_id, "Creating Embeddings", 3, 4, 0, len(all_chunks), "Generating vector embeddings...")
            
            # Add to vector store
            if all_chunks:
                logger.info(f"[STAGE 4/4] Adding {len(all_chunks)} chunks to vector store...")
                
                # Process in batches to avoid timeout
                batch_size = 50
                total_batches = (len(all_chunks) + batch_size - 1) // batch_size
                
                for i in range(0, len(all_chunks), batch_size):
                    batch = all_chunks[i:i+batch_size]
                    batch_num = (i // batch_size) + 1
                    logger.info(f"[EMBEDDING] Processing batch {batch_num}/{total_batches} ({len(batch)} chunks)")
                    
                    if progress_callback:
                        progress_callback(data_source_id, "Creating Embeddings", 3, 4, i, len(all_chunks), f"Embedding batch {batch_num}/{total_batches}...")
                    
                    success = await self.vector_service.add_documents(batch)
                    if not success:
                        logger.error(f"Failed to add batch {batch_num} to vector store")
                    else:
                        logger.info(f"[EMBEDDING] Batch {batch_num}/{total_batches} complete")
            
            logger.info(f"[COMPLETE] All documents processed and embedded successfully")
            db.commit()
            
        except Exception as e:
            db.rollback()
            logger.error(f"Error processing documents: {e}")
            raise
